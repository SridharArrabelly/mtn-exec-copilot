"""Create (or update) the Azure AI Search index used by the MTN Foundry agent.

Reads .docx, .pdf, .md, .txt files from ``data/`` (recursively), chunks them, embeds each chunk with the
embedding model deployed on the Foundry project (``text-embedding-3-small`` by
default), and uploads the chunks to an Azure AI Search index configured for
**hybrid search** (BM25 + vector) with a **semantic configuration** for
re-ranking.

Each chunk is enriched with date metadata extracted from the document filename
(pattern: ``Board Meeting – DD Month YYYY``). A ``meeting_date`` field enables
OData date-range filtering, and ``year``/``month`` fields support faceting. The
document title and date are prepended to every chunk so that embeddings always
carry temporal context — fixing issues where relative-date queries (e.g.
"February last year") retrieved the wrong meeting.

Embeddings are generated against the Foundry resource's Azure OpenAI route
(``/openai/deployments/<dep>/embeddings``), authenticated with
``DefaultAzureCredential`` — no separate ``AZURE_OPENAI_ENDPOINT`` required.

The vector dimension is **auto-detected** from the deployment at runtime, so
switching between ``text-embedding-3-small`` (1536) and ``text-embedding-3-large``
(3072) — or any other embedding model — is a one-env-var change. After switching
embedding models you MUST also set ``RECREATE_INDEX=true`` for one run, because
``vector_search_dimensions`` is immutable on an existing index.

Required environment variables (see ``.env.example``):
    AZURE_SEARCH_ENDPOINT      https://<svc>.search.windows.net
    SEARCH_INDEX_NAME          e.g. mtn-meetings
    PROJECT_ENDPOINT           https://<resource>.services.ai.azure.com/api/projects/<project>
    EMBEDDING_DEPLOYMENT       Foundry-deployed embedding model
                               (default: text-embedding-3-small)

Optional:
    AZURE_OPENAI_API_VERSION   default: 2024-10-21
    AZURE_SEARCH_API_KEY       if unset, uses DefaultAzureCredential
    DATA_DIR                   default: ./data
    CHUNK_SIZE                 chars per chunk, default: 1200
    CHUNK_OVERLAP              char overlap, default: 200
    RECREATE_INDEX             "true" to drop+recreate, default: false

Auth: ``az login``. Signed-in user needs:
  - "Search Index Data Contributor" + "Search Service Contributor" on the search service
  - "Azure AI User" (or equivalent) on the Foundry project

Usage:
    uv run python scripts/setup_aisearch_index.py
"""

from __future__ import annotations

import logging
import os
import re
import sys
import uuid
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional
from urllib.parse import urlparse

from openai import AzureOpenAI
from azure.core.credentials import AzureKeyCredential
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    AzureOpenAIVectorizer,
    AzureOpenAIVectorizerParameters,
    HnswAlgorithmConfiguration,
    HnswParameters,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SemanticConfiguration,
    SemanticField,
    SemanticPrioritizedFields,
    SemanticSearch,
    SimpleField,
    VectorSearch,
    VectorSearchAlgorithmMetric,
    VectorSearchProfile,
)
from docx import Document
from pypdf import PdfReader
from dotenv import load_dotenv

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)-7s %(message)s")
log = logging.getLogger("setup-aisearch")

EMBED_DIM_DEFAULT = 1536  # text-embedding-3-small; auto-detected at runtime
VECTOR_PROFILE = "mtn-vector-profile"
HNSW_ALGO = "mtn-hnsw"
SEMANTIC_CONFIG = "mtn-semantic"
VECTORIZER_NAME = "mtn-vectorizer"


# ---------- settings ----------

def _require(name: str) -> str:
    val = os.getenv(name, "").strip()
    if not val:
        log.error("Missing required env var: %s", name)
        sys.exit(1)
    return val


def load_settings() -> dict:
    return {
        "search_endpoint": _require("AZURE_SEARCH_ENDPOINT").rstrip("/"),
        "index_name": _require("SEARCH_INDEX_NAME"),
        "search_key": os.getenv("AZURE_SEARCH_API_KEY", "").strip(),
        "project_endpoint": _require("PROJECT_ENDPOINT").rstrip("/"),
        "embed_deployment": os.getenv("EMBEDDING_DEPLOYMENT", "text-embedding-3-small"),
        "aoai_api_version": os.getenv("AZURE_OPENAI_API_VERSION", "2024-10-21"),
        "data_dir": Path(os.getenv("DATA_DIR", "data")).resolve(),
        "chunk_size": int(os.getenv("CHUNK_SIZE", "1200")),
        "chunk_overlap": int(os.getenv("CHUNK_OVERLAP", "200")),
        "recreate": os.getenv("RECREATE_INDEX", "false").lower() == "true",
        # Filled in by detect_embed_dim() before ensure_index() runs.
        "embed_dim": None,
    }


# ---------- date extraction ----------

# Matches patterns like "Board Meeting – 15 February 2026" or "Board Meeting - 5 March 2019"
_DATE_RE = re.compile(
    r"(\d{1,2})\s+"
    r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+"
    r"(\d{4})",
    re.IGNORECASE,
)


def parse_meeting_date(filename: str) -> Optional[datetime]:
    """Extract meeting date from filename. Returns None if no date found."""
    m = _DATE_RE.search(filename)
    if not m:
        return None
    day, month_name, year = int(m.group(1)), m.group(2), int(m.group(3))
    try:
        return datetime(year, _month_number(month_name), day)
    except ValueError:
        return None


def _month_number(name: str) -> int:
    months = [
        "january", "february", "march", "april", "may", "june",
        "july", "august", "september", "october", "november", "december",
    ]
    return months.index(name.lower()) + 1


# ---------- clients ----------

def _aad():
    return DefaultAzureCredential()


def make_index_client(s: dict) -> SearchIndexClient:
    cred = AzureKeyCredential(s["search_key"]) if s["search_key"] else _aad()
    return SearchIndexClient(endpoint=s["search_endpoint"], credential=cred)


def make_search_client(s: dict) -> SearchClient:
    cred = AzureKeyCredential(s["search_key"]) if s["search_key"] else _aad()
    return SearchClient(endpoint=s["search_endpoint"], index_name=s["index_name"], credential=cred)


def make_embeddings_client(s: dict):
    """AzureOpenAI client pointed at the Foundry resource (for embeddings).

    Foundry's OpenAI-compat (/openai/v1/...) only serves chat/responses; the
    embeddings endpoint lives under the AOAI route /openai/deployments/<dep>/embeddings,
    which is reachable at the Foundry resource root.
    """
    parsed = urlparse(s["project_endpoint"])
    azure_endpoint = f"{parsed.scheme}://{parsed.netloc}"
    token_provider = get_bearer_token_provider(_aad(), "https://cognitiveservices.azure.com/.default")
    return AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_version=s["aoai_api_version"],
        azure_ad_token_provider=token_provider,
    )


# ---------- index ----------

def build_index(name: str, s: dict) -> SearchIndex:
    embed_dim = s.get("embed_dim") or EMBED_DIM_DEFAULT
    fields = [
        SimpleField(name="id", type=SearchFieldDataType.String, key=True, filterable=True),
        SearchableField(name="title", type=SearchFieldDataType.String, filterable=True, sortable=True),
        SimpleField(name="source", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="meeting_date", type=SearchFieldDataType.DateTimeOffset,
                    filterable=True, sortable=True, facetable=True),
        SimpleField(name="year", type=SearchFieldDataType.Int32,
                    filterable=True, sortable=True, facetable=True),
        SimpleField(name="month", type=SearchFieldDataType.Int32,
                    filterable=True, sortable=True, facetable=True),
        SimpleField(name="chunk_index", type=SearchFieldDataType.Int32, filterable=True, sortable=True),
        SearchableField(name="content", type=SearchFieldDataType.String, analyzer_name="en.microsoft"),
        SearchField(
            name="content_vector",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=embed_dim,
            vector_search_profile_name=VECTOR_PROFILE,
        ),
    ]

    # Build the Azure OpenAI vectorizer so AI Search can auto-vectorize text queries
    parsed = urlparse(s["project_endpoint"])
    azure_endpoint = f"{parsed.scheme}://{parsed.netloc}"

    vectorizer = AzureOpenAIVectorizer(
        vectorizer_name=VECTORIZER_NAME,
        parameters=AzureOpenAIVectorizerParameters(
            resource_url=azure_endpoint,
            deployment_name=s["embed_deployment"],
            model_name=s["embed_deployment"],
        ),
    )

    vector_search = VectorSearch(
        algorithms=[
            HnswAlgorithmConfiguration(
                name=HNSW_ALGO,
                parameters=HnswParameters(metric=VectorSearchAlgorithmMetric.COSINE),
            )
        ],
        profiles=[
            VectorSearchProfile(
                name=VECTOR_PROFILE,
                algorithm_configuration_name=HNSW_ALGO,
                vectorizer_name=VECTORIZER_NAME,
            )
        ],
        vectorizers=[vectorizer],
    )

    semantic_search = SemanticSearch(
        configurations=[
            SemanticConfiguration(
                name=SEMANTIC_CONFIG,
                prioritized_fields=SemanticPrioritizedFields(
                    title_field=SemanticField(field_name="title"),
                    content_fields=[SemanticField(field_name="content")],
                ),
            )
        ]
    )

    return SearchIndex(
        name=name,
        fields=fields,
        vector_search=vector_search,
        semantic_search=semantic_search,
    )


def ensure_index(s: dict) -> None:
    client = make_index_client(s)
    existing_indexes = {i.name: i for i in client.list_indexes()}

    # Vector dimensions are immutable on an existing index. If the user
    # switched embedding models without setting RECREATE_INDEX=true, the
    # downstream create_or_update_index call would fail with a confusing
    # server-side error — surface the real cause up front.
    if s["index_name"] in existing_indexes and not s["recreate"]:
        existing = existing_indexes[s["index_name"]]
        existing_dim = next(
            (
                getattr(f, "vector_search_dimensions", None)
                for f in existing.fields
                if f.name == "content_vector"
            ),
            None,
        )
        if existing_dim and s.get("embed_dim") and existing_dim != s["embed_dim"]:
            log.error(
                "Existing index '%s' has vector dim %d but the deployment '%s' produces %d-dim embeddings. "
                "Vector dimensions are immutable on an existing index — set RECREATE_INDEX=true to drop "
                "and rebuild from scratch.",
                s["index_name"], existing_dim, s["embed_deployment"], s["embed_dim"],
            )
            sys.exit(2)

    if s["index_name"] in existing_indexes and s["recreate"]:
        log.info("Deleting existing index '%s'", s["index_name"])
        client.delete_index(s["index_name"])
        existing_indexes.pop(s["index_name"], None)

    if s["index_name"] in existing_indexes:
        log.info("Updating index '%s'", s["index_name"])
        client.create_or_update_index(build_index(s["index_name"], s))
    else:
        log.info("Creating index '%s'", s["index_name"])
        client.create_index(build_index(s["index_name"], s))


# ---------- ingest ----------

def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "").strip() for page in reader.pages)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


# Map file extension -> reader. Extend here to support more formats.
READERS = {
    ".docx": read_docx,
    ".pdf": read_pdf,
    ".md": read_text,
    ".markdown": read_text,
    ".txt": read_text,
}


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    text = re.sub(r"[ \t]+", " ", text).strip()
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + size, n)
        if end < n:
            window = text[start:end]
            for sep in ("\n\n", "\n", ". ", " "):
                idx = window.rfind(sep)
                if idx >= int(size * 0.6):
                    end = start + idx + len(sep)
                    break
        chunks.append(text[start:end].strip())
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return [c for c in chunks if c]


def embed_batch(client, deployment: str, texts: list[str]) -> list[list[float]]:
    resp = client.embeddings.create(model=deployment, input=texts)
    return [d.embedding for d in resp.data]


def detect_embed_dim(client, deployment: str) -> int:
    """Probe the deployment to learn its embedding output dimension.

    Lets the script work with any embedding model (3-small=1536, 3-large=3072,
    ada-002=1536, etc.) without hardcoding. Costs one tiny embedding call at
    setup time.
    """
    resp = client.embeddings.create(model=deployment, input=["dimension probe"])
    return len(resp.data[0].embedding)


def iter_documents(s: dict, aoai) -> Iterable[dict]:
    files = sorted(
        f for f in s["data_dir"].rglob("*")
        if f.is_file() and f.suffix.lower() in READERS
    )
    if not files:
        log.warning("No supported files (%s) found in %s",
                    ", ".join(sorted(READERS)), s["data_dir"])
        return

    for path in files:
        title = path.stem
        reader = READERS[path.suffix.lower()]
        log.info("Reading %s", path.name)
        try:
            raw = reader(path)
        except Exception as e:
            log.warning("  failed to read %s: %s", path.name, e)
            continue

        # Extract meeting date from filename
        meeting_dt = parse_meeting_date(path.stem)
        if meeting_dt:
            date_prefix = f"[Document: {title} | Meeting Date: {meeting_dt.strftime('%d %B %Y')}]\n\n"
            log.info("  meeting date: %s", meeting_dt.strftime("%Y-%m-%d"))
        else:
            date_prefix = f"[Document: {title}]\n\n"
            log.warning("  no date found in filename")

        chunks = chunk_text(raw, s["chunk_size"], s["chunk_overlap"])
        if not chunks:
            log.warning("  no text extracted from %s", path.name)
            continue
        log.info("  %d chunks", len(chunks))

        # Prepend date/title context to each chunk for embedding
        enriched_chunks = [date_prefix + chunk for chunk in chunks]

        BATCH = 16
        for i in range(0, len(enriched_chunks), BATCH):
            batch = enriched_chunks[i : i + BATCH]
            vectors = embed_batch(aoai, s["embed_deployment"], batch)
            for j, (text, vec) in enumerate(zip(batch, vectors)):
                idx = i + j
                doc = {
                    "id": f"{uuid.uuid5(uuid.NAMESPACE_URL, f'{path.name}:{idx}')}",
                    "title": title,
                    "source": path.name,
                    "chunk_index": idx,
                    "content": text,
                    "content_vector": vec,
                }
                if meeting_dt:
                    doc["meeting_date"] = meeting_dt.isoformat() + "Z"
                    doc["year"] = meeting_dt.year
                    doc["month"] = meeting_dt.month
                yield doc


def upload(s: dict, docs: Iterable[dict]) -> int:
    search = make_search_client(s)
    BATCH = 100
    buf: list[dict] = []
    total = 0
    for d in docs:
        buf.append(d)
        if len(buf) >= BATCH:
            search.upload_documents(documents=buf)
            total += len(buf)
            log.info("  uploaded %d (running total %d)", len(buf), total)
            buf.clear()
    if buf:
        search.upload_documents(documents=buf)
        total += len(buf)
        log.info("  uploaded %d (running total %d)", len(buf), total)
    return total


# ---------- main ----------

def main() -> None:
    s = load_settings()
    log.info("Search:    %s  /  index=%s", s["search_endpoint"], s["index_name"])
    log.info("Foundry:   %s  /  embed=%s", s["project_endpoint"], s["embed_deployment"])
    log.info("Data dir:  %s", s["data_dir"])

    aoai = make_embeddings_client(s)
    s["embed_dim"] = detect_embed_dim(aoai, s["embed_deployment"])
    log.info("Embedding dim (detected): %d", s["embed_dim"])

    ensure_index(s)
    n = upload(s, iter_documents(s, aoai))
    log.info("Done. Indexed %d chunks into '%s'.", n, s["index_name"])


if __name__ == "__main__":
    main()
